from flask import Flask, render_template, request, jsonify, send_file
import requests
import io
import json
from docx import Document

import tempfile

app = Flask(__name__)

@app.route("/project/<int:project_id>/report-docx")
def project_report_docx(project_id):
    api_key = request.args.get("api_key")
    base_url = request.args.get("base_url")

    # Fetch data
    project = requests.get(f"{base_url}/v1/project/view", params={"id": project_id, "api_key": api_key}).json()
    mistakes = requests.get(f"{base_url}/v1/project/mistake-counts", params={"id": project_id, "api_key": api_key}).json()

    # Create document
    doc = Document()
    doc.add_heading("Project Report", 0)

    doc.add_heading("Main", level=1)
    doc.add_paragraph(f"ID: {project['id']}")
    doc.add_paragraph(f"Client: {project['client']['name']}")
    doc.add_paragraph(f"Name: {project['name']}")
    doc.add_paragraph(f"Spec: {project['spec']['name']}")
    doc.add_paragraph(f"Service: {project['service']['name']}")
    doc.add_paragraph(f"Status: {stage_map(project['stage_id'])}")
    doc.add_paragraph(f"Evaluation Summary: {project['evaluation_summary']}")
    doc.add_paragraph(f"Score: {project['score']}")
    doc.add_paragraph(f"Mark: {project['mark']}")
    doc.add_paragraph(f"Created at: {project['created_at']}")

    doc.add_heading("Languages", level=1)
    doc.add_paragraph(f"Source: {project['sourceLang']['tag']}")
    doc.add_paragraph(f"Target: {project['targetLang']['tag']}")

    doc.add_heading("Project Team", level=1)
    doc.add_paragraph(f"Manager: {format_person(project['manager'])}")
    doc.add_paragraph(f"Translator: {format_person(project['translator'])}")
    doc.add_paragraph(f"Evaluator: {format_person(project['evaluator'])}")
    doc.add_paragraph(f"Arbiter: {format_person(project['arbiter'])}")

    doc.add_heading("Evaluation Details", level=1)
    doc.add_paragraph(f"Word count: {project['word_count_for_evaluator'] or '–'}")
    doc.add_paragraph(f"Note: {project['note_for_evaluator'] or '–'}")
    doc.add_paragraph(f"Evaluation Count: {project['evaluation_count']}")

    doc.add_heading("Mistake Breakdown", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Type'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'Count'

    for m in mistakes:
        row_cells = table.add_row().cells
        row_cells[0].text = m['type']['name']
        row_cells[1].text = m['severity']['name']
        row_cells[2].text = str(m['count'])

    # Save to a temporary file
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.seek(0)

    return send_file(tmp.name, as_attachment=True, download_name=f"Project_{project_id}_Report.docx")

@app.route("/", methods=["GET", "POST"])
def index():
    projects = []
    error = None
    api_key = ""
    base_url = ""

    if request.method == "POST":
        api_key = request.form.get("api_key")
        base_url = request.form.get("base_url").rstrip("/")

        try:
            response = requests.get(f"{base_url}/v1/project/index", params={"api_key": api_key})
            response.raise_for_status()
            projects = response.json()
        except requests.exceptions.HTTPError as e:
            error = f"HTTP error: {e.response.status_code} - {e.response.json().get('detail', '')}"
        except Exception as e:
            error = f"Unexpected error: {str(e)}"

    return render_template("index.html", projects=projects, error=error, api_key=api_key, base_url=base_url)


@app.route("/project/<int:project_id>/details")
def project_details(project_id):
    api_key = request.args.get("api_key")
    base_url = request.args.get("base_url")
    try:
        response = requests.get(f"{base_url}/v1/project/view", params={"id": project_id, "api_key": api_key})
        response.raise_for_status()
        return jsonify(response.json())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/project/<int:project_id>/mistakes")
def project_mistakes(project_id):
    api_key = request.args.get("api_key")
    base_url = request.args.get("base_url")
    try:
        response = requests.get(f"{base_url}/v1/project/mistake-counts", params={"id": project_id, "api_key": api_key})
        response.raise_for_status()
        json_data = response.json()
        return send_file(
            io.BytesIO(json.dumps(json_data, indent=2).encode("utf-8")),
            mimetype='application/json',
            as_attachment=True,
            download_name=f"mistakes_project_{project_id}.json"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def format_person(p):
    return f"{p['first_name']} {p['last_name']}, {p['company']}"

def stage_map(stage_id):
    return {
        1: "Files upload",
        2: "Comparison",
        3: "Evaluation",
        4: "Translator's review",
        5: "Arbitration",
        6: "Completed"
    }.get(stage_id, "Unknown")


if __name__ == "__main__":
    app.run(port=5000)
